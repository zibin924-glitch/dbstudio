"""Document generator - produces Markdown, DOCX, and PDF documentation from database metadata."""

import logging
from typing import Any

logger = logging.getLogger(__name__)

# Attempt to import WeasyPrint; it requires system libraries (cairo, pango).
_WEASYPRINT_AVAILABLE = False
try:
    import weasyprint  # type: ignore[import-untyped]
    _WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as exc:
    logger.info(
        "WeasyPrint is not available (%s). PDF output will fall back to HTML.",
        exc,
    )


class DocGenerator:
    """Generates structured database documentation in multiple formats."""

    @staticmethod
    def generate_markdown(
        tables: list[dict],
        db_name: str,
        db_type: str,
    ) -> str:
        """Generate a structured Markdown document from table metadata.

        Args:
            tables: List of table dicts, each containing keys:
                    name, columns, indexes, foreign_keys, comment, ddl (optional).
            db_name: Database name for the document title.
            db_type: Database type (mysql, postgresql, oracle).

        Returns:
            Complete Markdown string with table of contents, table details,
            columns, indexes, foreign keys, and DDL.
        """
        lines: list[str] = []

        # Title
        lines.append(f"# Database Documentation: {db_name}")
        lines.append("")
        lines.append(f"**Database Type:** {db_type}")
        lines.append(f"**Total Tables:** {len(tables)}")
        lines.append("")

        # Table of contents
        lines.append("## Table of Contents")
        lines.append("")
        for i, table in enumerate(tables, start=1):
            name = table.get("name", f"table_{i}")
            anchor = name.lower().replace(" ", "-")
            comment = table.get("comment", "")
            desc = f" - {comment}" if comment else ""
            lines.append(f"{i}. [{name}](#{anchor}){desc}")
        lines.append("")
        lines.append("---")
        lines.append("")

        # Table details
        for table in tables:
            name = table.get("name", "Unknown")
            comment = table.get("comment", "")
            columns = table.get("columns", [])
            indexes = table.get("indexes", [])
            foreign_keys = table.get("foreign_keys", [])
            ddl = table.get("ddl", "")

            lines.append(f"## {name}")
            lines.append("")
            if comment:
                lines.append(f"> {comment}")
                lines.append("")

            # Columns table
            if columns:
                lines.append("### Columns")
                lines.append("")
                lines.append("| # | Name | Type | Nullable | Default | Primary Key | Auto Increment | Comment |")
                lines.append("|---|------|------|----------|---------|-------------|----------------|---------|")
                for idx, col in enumerate(columns, start=1):
                    pk = "Yes" if col.get("primary_key") else ""
                    ai = "Yes" if col.get("auto_increment") else ""
                    nullable = "Yes" if col.get("nullable", True) else "No"
                    default = col.get("default") or ""
                    col_comment = col.get("comment") or ""
                    lines.append(
                        f"| {idx} | {col['name']} | {col['type']} | {nullable} | "
                        f"{default} | {pk} | {ai} | {col_comment} |"
                    )
                lines.append("")

            # Indexes
            if indexes:
                lines.append("### Indexes")
                lines.append("")
                lines.append("| Name | Type | Columns | Unique |")
                lines.append("|------|------|---------|--------|")
                for idx_info in indexes:
                    cols = ", ".join(idx_info.get("columns", []))
                    unique = "Yes" if idx_info.get("unique") else "No"
                    lines.append(
                        f"| {idx_info.get('name', '')} | {idx_info.get('type', '')} | "
                        f"{cols} | {unique} |"
                    )
                lines.append("")

            # Foreign keys
            if foreign_keys:
                lines.append("### Foreign Keys")
                lines.append("")
                lines.append("| Name | Source Column | Target Table | Target Column | On Update | On Delete |")
                lines.append("|------|---------------|--------------|---------------|-----------|-----------|")
                for fk in foreign_keys:
                    lines.append(
                        f"| {fk.get('name', '')} | {fk.get('source_column', '')} | "
                        f"{fk.get('target_table', '')} | {fk.get('target_column', '')} | "
                        f"{fk.get('on_update', 'NO ACTION')} | {fk.get('on_delete', 'NO ACTION')} |"
                    )
                lines.append("")

            # DDL
            if ddl:
                lines.append("### DDL")
                lines.append("")
                lines.append("```sql")
                lines.append(ddl)
                lines.append("```")
                lines.append("")

            lines.append("---")
            lines.append("")

        return "\n".join(lines)

    @staticmethod
    def generate_docx(
        tables: list[dict],
        db_name: str,
        db_type: str,
        output_path: str,
    ) -> None:
        """Generate a Word document (.docx) from table metadata.

        Args:
            tables: List of table metadata dicts.
            db_name: Database name.
            db_type: Database type string.
            output_path: File path to write the .docx to.
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f"Database Documentation: {db_name}", level=0)

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Database Type: {db_type}\n").bold = True
        meta_para.add_run(f"Total Tables: {len(tables)}\n")

        doc.add_page_break()

        # Table of Contents header
        doc.add_heading("Table of Contents", level=1)
        for i, table in enumerate(tables, start=1):
            name = table.get("name", f"table_{i}")
            comment = table.get("comment", "")
            desc = f" - {comment}" if comment else ""
            doc.add_paragraph(f"{i}. {name}{desc}", style="List Number")

        doc.add_page_break()

        # Table details
        for table in tables:
            name = table.get("name", "Unknown")
            comment = table.get("comment", "")
            columns = table.get("columns", [])
            indexes = table.get("indexes", [])
            foreign_keys = table.get("foreign_keys", [])

            doc.add_heading(name, level=2)
            if comment:
                doc.add_paragraph(comment, style="Intense Quote")

            # Columns table
            if columns:
                doc.add_heading("Columns", level=3)
                col_table = doc.add_table(rows=1, cols=7)
                col_table.style = "Light Grid Accent 1"

                # Header row
                header_cells = col_table.rows[0].cells
                headers = ["Name", "Type", "Nullable", "Default", "PK", "Auto Inc", "Comment"]
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Data rows
                for col in columns:
                    row_cells = col_table.add_row().cells
                    row_cells[0].text = col.get("name", "")
                    row_cells[1].text = col.get("type", "")
                    row_cells[2].text = "Yes" if col.get("nullable", True) else "No"
                    row_cells[3].text = str(col.get("default", "") or "")
                    row_cells[4].text = "Yes" if col.get("primary_key") else ""
                    row_cells[5].text = "Yes" if col.get("auto_increment") else ""
                    row_cells[6].text = col.get("comment", "") or ""

            # Indexes
            if indexes:
                doc.add_heading("Indexes", level=3)
                idx_table = doc.add_table(rows=1, cols=4)
                idx_table.style = "Light Grid Accent 1"
                header_cells = idx_table.rows[0].cells
                for idx, header in enumerate(["Name", "Type", "Columns", "Unique"]):
                    header_cells[idx].text = header
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for idx_info in indexes:
                    row_cells = idx_table.add_row().cells
                    row_cells[0].text = idx_info.get("name", "")
                    row_cells[1].text = idx_info.get("type", "")
                    row_cells[2].text = ", ".join(idx_info.get("columns", []))
                    row_cells[3].text = "Yes" if idx_info.get("unique") else "No"

            # Foreign Keys
            if foreign_keys:
                doc.add_heading("Foreign Keys", level=3)
                fk_table = doc.add_table(rows=1, cols=6)
                fk_table.style = "Light Grid Accent 1"
                header_cells = fk_table.rows[0].cells
                for idx, header in enumerate([
                    "Name", "Source Column", "Target Table",
                    "Target Column", "On Update", "On Delete",
                ]):
                    header_cells[idx].text = header
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                for fk in foreign_keys:
                    row_cells = fk_table.add_row().cells
                    row_cells[0].text = fk.get("name", "")
                    row_cells[1].text = fk.get("source_column", "")
                    row_cells[2].text = fk.get("target_table", "")
                    row_cells[3].text = fk.get("target_column", "")
                    row_cells[4].text = fk.get("on_update", "NO ACTION")
                    row_cells[5].text = fk.get("on_delete", "NO ACTION")

            doc.add_page_break()

        doc.save(output_path)
        logger.info("DOCX document saved to %s", output_path)

    @staticmethod
    def _build_html_from_markdown(markdown_content: str, db_name: str) -> str:
        """Convert markdown content to a styled HTML string suitable for PDF or browser display.

        Args:
            markdown_content: Markdown text generated by generate_markdown().
            db_name: Database name used in the HTML <title>.

        Returns:
            A complete HTML document string.
        """
        html_parts: list[str] = []
        html_parts.append("<!DOCTYPE html>")
        html_parts.append("<html><head>")
        html_parts.append(f"<title>Database Documentation: {db_name}</title>")
        html_parts.append("""
        <style>
            body { font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; color: #333; }
            h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }
            h2 { color: #2980b9; margin-top: 30px; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }
            h3 { color: #7f8c8d; }
            table { border-collapse: collapse; width: 100%; margin: 10px 0 20px 0; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; font-size: 13px; }
            th { background-color: #3498db; color: white; }
            tr:nth-child(even) { background-color: #f2f2f2; }
            code { background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }
            pre { background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }
            blockquote { border-left: 4px solid #3498db; margin-left: 0; padding-left: 15px; color: #666; }
            @media print { body { margin: 20px; } }
        </style>
        """)
        html_parts.append("</head><body>")

        # Simple markdown-to-HTML conversion
        in_code_block = False
        in_table = False
        table_header_done = False

        for line in markdown_content.split("\n"):
            stripped = line.strip()

            # Code blocks
            if stripped.startswith("```"):
                if in_code_block:
                    html_parts.append("</code></pre>")
                    in_code_block = False
                else:
                    html_parts.append("<pre><code>")
                    in_code_block = True
                continue

            if in_code_block:
                # Escape HTML entities in code
                escaped = (
                    stripped.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                html_parts.append(escaped + "\n")
                continue

            # Table rows
            if stripped.startswith("|"):
                if not in_table:
                    html_parts.append("<table>")
                    in_table = True
                    table_header_done = False

                # Skip separator rows
                if set(stripped.replace("|", "").replace("-", "").replace(" ", "")) == set():
                    table_header_done = True
                    continue

                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                if not table_header_done:
                    tag = "th"
                else:
                    tag = "td"

                row_html = "<tr>" + "".join(f"<{tag}>{c}</{tag}>" for c in cells) + "</tr>"
                html_parts.append(row_html)
                continue
            elif in_table:
                html_parts.append("</table>")
                in_table = False
                table_header_done = False

            # Headings
            if stripped.startswith("# "):
                html_parts.append(f"<h1>{stripped[2:]}</h1>")
            elif stripped.startswith("## "):
                html_parts.append(f"<h2>{stripped[3:]}</h2>")
            elif stripped.startswith("### "):
                html_parts.append(f"<h3>{stripped[4:]}</h3>")
            elif stripped.startswith("> "):
                html_parts.append(f"<blockquote>{stripped[2:]}</blockquote>")
            elif stripped.startswith("---"):
                html_parts.append("<hr>")
            elif stripped.startswith("**") and stripped.endswith("**"):
                html_parts.append(f"<p><strong>{stripped[2:-2]}</strong></p>")
            elif stripped.startswith("1. ") or stripped.startswith("- "):
                text = stripped[3:] if stripped.startswith("1. ") else stripped[2:]
                html_parts.append(f"<p>{text}</p>")
            elif stripped:
                html_parts.append(f"<p>{stripped}</p>")

        if in_table:
            html_parts.append("</table>")
        if in_code_block:
            html_parts.append("</code></pre>")

        html_parts.append("</body></html>")
        return "\n".join(html_parts)

    @staticmethod
    def generate_pdf(
        tables: list[dict],
        db_name: str,
        db_type: str,
        output_path: str,
    ) -> str:
        """Generate a PDF document from table metadata using WeasyPrint.

        Converts the Markdown documentation to styled HTML and then uses
        WeasyPrint to render it as a real PDF. If WeasyPrint (or its
        system dependencies cairo/pango) is not installed, the HTML is
        written directly as a graceful fallback.

        Args:
            tables: List of table metadata dicts.
            db_name: Database name.
            db_type: Database type string.
            output_path: File path to write the output to.

        Returns:
            A string indicating the actual format written:
            ``"pdf"`` if WeasyPrint succeeded, ``"html"`` on fallback.
        """
        # Generate markdown first
        markdown_content = DocGenerator.generate_markdown(tables, db_name, db_type)
        html_content = DocGenerator._build_html_from_markdown(markdown_content, db_name)

        # Attempt real PDF generation via WeasyPrint
        if _WEASYPRINT_AVAILABLE:
            try:
                pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
                with open(output_path, "wb") as f:
                    f.write(pdf_bytes)
                logger.info("PDF document saved to %s (WeasyPrint)", output_path)
                return "pdf"
            except Exception as exc:
                logger.warning(
                    "WeasyPrint PDF generation failed (%s). Falling back to HTML.",
                    exc,
                )

        # Fallback: write the HTML directly
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        logger.info("PDF fallback (HTML) document saved to %s", output_path)
        return "html"

    @staticmethod
    def generate_pdf_bytes(
        tables: list[dict],
        db_name: str,
        db_type: str,
    ) -> tuple[bytes, str]:
        """Generate PDF (or HTML fallback) as in-memory bytes.

        This is a convenience wrapper around :meth:`generate_pdf` that
        avoids writing to disk, making it suitable for streaming HTTP
        responses.

        Args:
            tables: List of table metadata dicts.
            db_name: Database name.
            db_type: Database type string.

        Returns:
            A ``(content_bytes, format_string)`` tuple where
            *format_string* is ``"pdf"`` or ``"html"``.
        """
        markdown_content = DocGenerator.generate_markdown(tables, db_name, db_type)
        html_content = DocGenerator._build_html_from_markdown(markdown_content, db_name)

        if _WEASYPRINT_AVAILABLE:
            try:
                pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
                logger.info("PDF document generated in memory (WeasyPrint)")
                return pdf_bytes, "pdf"
            except Exception as exc:
                logger.warning(
                    "WeasyPrint PDF generation failed (%s). Falling back to HTML.",
                    exc,
                )

        return html_content.encode("utf-8"), "html"
